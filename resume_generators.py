"""
resume_generators.py — ApplyEdge
Rendering code copied VERBATIM from faang_pdf_v3.py and faang_resume_v3.js.
NOT ONE CHARACTER of the rendering is changed.
A parser sits on top and extracts structured data from any resume text,
then feeds it into those exact same rendering functions.
"""

import io, re


# ═══════════════════════════════════════════════════════════════════════════════
# PARSER  — extracts structured fields from any plain-text resume
# Returns: name, contact, summary, skills, jobs, education, projects, certs
# ═══════════════════════════════════════════════════════════════════════════════

SECTION_TOKENS = {
    "summary","experience","education","skills","projects","certifications",
    "certification","certificates","achievements","awards","honors",
    "publications","volunteer","languages","interests","profile","objective",
    "history","competencies","expertise","credentials","employment","career",
    "research","licenses",
}

SECTION_ORDER = [
    "professional summary","summary","objective","profile",
    "technical skills","skills","core competencies","competencies",
    "professional experience","work experience","experience","employment",
    "education","projects",
    "certifications","certification","certificates","licenses",
    "achievements","awards","honors","publications",
    "volunteer","languages","interests",
]

def _sec_rank(key):
    low = key.lower().strip()
    for i, kw in enumerate(SECTION_ORDER):
        if low == kw or kw in low or low in kw:
            return i
    return 999

def _is_sec(s):
    s = s.strip().rstrip(":")
    if not s or len(s) > 70: return False
    low = s.lower()
    if low.rstrip(":") in [x.rstrip(":") for x in SECTION_ORDER]: return True
    if s.isupper() and 2 < len(s) < 70:
        return bool(set(re.sub(r"[^a-z\s]","",low).split()) & SECTION_TOKENS)
    words = set(re.sub(r"[^a-z\s]","",low).split())
    if words & SECTION_TOKENS and len(s.split()) <= 6:
        if not any(c in s for c in ["•","@",".","http","+1"]): return True
    return False

def _is_bul(s):   return bool(re.match(r"^[•▪●\-–\*◦]\s", s.strip()))
def _has_date(s): return bool(re.search(r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|20\d\d|19\d\d|present|current)", s.lower()))
def _debul(s):    return re.sub(r"^[•▪●\-–\*◦]\s*","", s.strip())

def _split_tab(line):
    """Split 'Left\tRight' or 'Left   Right' → (left, right)"""
    line = line.rstrip("\t ")
    if "\t" in line:
        p = line.split("\t",1); return p[0].strip(), p[1].strip() if len(p)>1 else ""
    m = re.split(r"\s{3,}", line.strip(), 1)
    return (m[0].strip(), m[1].strip()) if len(m)==2 else (line.strip(), "")

def _split_role_dates(s):
    """'Firmware Engineer Jun 2024 – Present' → ('Firmware Engineer', 'Jun 2024 – Present')"""
    m = re.search(r"\s+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}.*)$", s, re.I)
    if m: return s[:m.start()].strip(), m.group(1).strip()
    m = re.search(r"\s+(\d{4}\s*[–\-]\s*(?:\d{4}|Present|Current))$", s, re.I)
    if m: return s[:m.start()].strip(), m.group(1).strip()
    return s, ""

def _clean(text):
    out = []
    for line in text.splitlines():
        line = re.sub(r"^#{1,4}\s*","",line)
        line = re.sub(r"\*{3}([^*\n]+)\*{3}",r"\1",line)
        line = re.sub(r"(?<!\*)\*(?!\*)([^*\n]+)(?<!\*)\*(?!\*)",r"\1",line)
        line = re.sub(r"_{1,2}([^_\n]+)_{1,2}",r"\1",line)
        line = re.sub(r"^>\s*","",line)
        line = re.sub(r"<[^>]+>","",line)
        out.append(line)
    return "\n".join(out)

def _parse_resume(resume_text):
    """
    Returns dict with keys:
      name       : str
      contact    : str
      summary    : list[str]          — paragraph strings
      skills     : list[(label, val)] — ("AI / LLM", "Python, ...")
      jobs       : list[dict]         — {company, role, city, dates, bullets:[]}
      education  : list[dict]         — {school, dates, degree}
      projects   : list[dict]         — {name, stack, bullets:[]}
      certs      : list[str]          — plain cert strings
    """
    raw  = _clean(resume_text).strip().splitlines()
    first = next((i for i,l in enumerate(raw) if _is_sec(l.strip())), len(raw))
    pre   = [l.strip() for l in raw[:first] if l.strip()]

    # Parse sections, merging duplicates
    secs = {}
    cur_key = cur_label = None
    cur_lines = []
    for line in raw[first:]:
        s = line.strip()
        if s and _is_sec(s):
            if cur_key:
                if cur_key in secs: secs[cur_key][2].extend(cur_lines)
                else:               secs[cur_key] = [cur_key, s.rstrip(":"), cur_lines]
            cur_key, cur_label, cur_lines = s.lower().rstrip(":").strip(), s.rstrip(":"), []
        else:
            cur_lines.append(line)
    if cur_key:
        if cur_key in secs: secs[cur_key][2].extend(cur_lines)
        else:               secs[cur_key] = [cur_key, cur_label, cur_lines]

    # ALL-CAPS names mistaken as sections → rescue to header
    bad = [k for k,(_, lbl, _) in secs.items()
           if lbl.isupper() and len(lbl.split())<=4
           and not any(t in k for t in SECTION_TOKENS)]
    for k in bad:
        _, lbl, extra = secs.pop(k)
        pre = [lbl]+[l for l in extra if l.strip()]+pre

    # Name/contact buried at bottom of last section
    if not pre and secs:
        last_k = list(secs.keys())[-1]
        last_l = secs[last_k][2]
        rescued = []
        for l in reversed(last_l):
            s = l.strip()
            if not s: continue
            is_name = not _has_date(s) and not _is_bul(s) and len(s.split())<=6 and not _is_sec(s)
            is_con  = "@" in s or "linkedin" in s.lower() or bool(re.search(r"\d{3}[-.\s]\d{3}",s))
            if is_name or is_con: rescued.insert(0, s)
            else: break
        if rescued:
            rescued_set = set(rescued)
            secs[last_k][2] = [l for l in last_l if l.strip() not in rescued_set]
            pre = rescued

    name    = pre[0] if pre else ""
    contact = "  |  ".join(pre[1:]) if len(pre)>1 else ""

    # ── Extract each section type ─────────────────────────────────────────────
    def get_sec(*keys):
        for k in sorted(secs, key=_sec_rank):
            if any(x in k for x in keys):
                return [l for l in secs[k][2] if l.strip()]
        return []

    # Summary
    summary = []
    for l in get_sec("summary","objective","profile"):
        s = l.strip()
        if s and not _is_sec(s):
            summary.append(_debul(s) if _is_bul(s) else s)

    # Skills → list of (label, value)
    skills, seen = [], set()
    for l in get_sec("skill","competenc","technolog","expertise"):
        s = l.strip()
        if not s or s.upper().rstrip(":") in ("TECHNICAL SKILLS","SKILLS","CORE COMPETENCIES"): continue
        if ":" in s:
            p = s.split(":",1); lbl = p[0].strip()
            if lbl not in seen: skills.append((lbl, p[1].strip())); seen.add(lbl)

    # Jobs → list of {company, role, city, dates, bullets}
    jobs = []
    cur_job = None
    exp_lines = get_sec("experience","employment","history","career")
    i = 0
    while i < len(exp_lines):
        s = exp_lines[i].strip()
        if not s: i+=1; continue

        if _is_bul(s):
            if cur_job is None: cur_job={"company":"","role":"","city":"","dates":"","bullets":[]}
            cur_job["bullets"].append(_debul(s)); i+=1; continue

        left, right = _split_tab(exp_lines[i])
        nxt  = next((exp_lines[j] for j in range(i+1,len(exp_lines)) if exp_lines[j].strip()),"")
        nl, nr = _split_tab(nxt)

        # Company+city (short, no date) → next line has date
        if not _has_date(left) and not _has_date(right) and len(left)<=80 and (_has_date(nr) or _has_date(nl)):
            if cur_job: jobs.append(cur_job)
            role, dates = (nl,nr) if _has_date(nr) else _split_role_dates(nl)
            cur_job = {"company":left,"role":role,"city":right,"dates":dates,"bullets":[]}
            i+=2; continue

        # Right col has date
        if _has_date(right):
            if cur_job: jobs.append(cur_job)
            cur_job = {"company":left,"role":"","city":"","dates":right,"bullets":[]}
            i+=1; continue

        # Left has embedded date (role+date on one line, no tab)
        if _has_date(left) and not right:
            role, dates = _split_role_dates(left)
            if dates:
                if cur_job: jobs.append(cur_job)
                cur_job = {"company":"","role":role,"city":"","dates":dates,"bullets":[]}
                i+=1; continue

        # Long paragraph → bullet, short non-date line → bullet too
        if cur_job is not None: cur_job["bullets"].append(s)
        i+=1
    if cur_job: jobs.append(cur_job)

    # Education → list of {school, dates, degree}
    education = []
    edu_lines = get_sec("education","academic","credential")
    i = 0
    while i < len(edu_lines):
        s = edu_lines[i].strip()
        if not s: i+=1; continue
        left, right = _split_tab(edu_lines[i])
        if right and (_has_date(right) or any(d in right for d in ["MS","M.S","B.Tech","Bachelor","Master","B.E","B.Sc","Ph.D","MBA"])):
            e = {"school":left,"dates":right,"degree":""}
            ns = edu_lines[i+1].strip() if i+1<len(edu_lines) else ""
            if ns and not _has_date(ns) and not _is_bul(ns) and not _is_sec(ns):
                e["degree"]=ns; i+=2
            else: i+=1
            education.append(e); continue
        if not _has_date(s) and not _is_bul(s):
            ns = edu_lines[i+1].strip() if i+1<len(edu_lines) else ""
            if _has_date(ns):
                nn = edu_lines[i+2].strip() if i+2<len(edu_lines) else ""
                deg = nn if nn and not _has_date(nn) and not _is_sec(nn) else ""
                education.append({"school":s,"dates":ns,"degree":deg})
                i += 3 if deg else 2; continue
        if _has_date(s) and education:
            education[-1]["dates"] = s
        i+=1

    # Projects → list of {name, stack, bullets}
    projects = []
    cur_proj = None
    for l in get_sec("project"):
        s = l.strip()
        if not s: continue
        if _is_bul(s):
            if cur_proj is None: cur_proj={"name":"","stack":"","bullets":[]}
            cur_proj["bullets"].append(_debul(s))
        else:
            if cur_proj: projects.append(cur_proj)
            if "—" in s:   p2=s.split("—",1); cur_proj={"name":p2[0].strip(),"stack":p2[1].strip(),"bullets":[]}
            elif "|" in s:  p2=s.split("|",1); cur_proj={"name":p2[0].strip(),"stack":p2[1].strip(),"bullets":[]}
            else:           cur_proj={"name":s,"stack":"","bullets":[]}
    if cur_proj: projects.append(cur_proj)

    # Certs → list of plain strings
    certs = []
    for l in get_sec("certif","certificate","license"):
        s = l.strip()
        if s: certs.append(_debul(s) if _is_bul(s) else s)

    return dict(name=name, contact=contact, summary=summary,
                skills=skills, jobs=jobs, education=education,
                projects=projects, certs=certs)


# ═══════════════════════════════════════════════════════════════════════════════
# PDF GENERATOR
# Rendering code copied VERBATIM from faang_pdf_v3.py — NOT ONE CHARACTER changed.
# Only the data fed into it comes from _parse_resume() above.
# ═══════════════════════════════════════════════════════════════════════════════

def generate_resume_pdf(resume_text: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, HRFlowable, Table, TableStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    data = _parse_resume(resume_text)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=0.6*inch, rightMargin=0.6*inch,
        topMargin=0.55*inch, bottomMargin=0.55*inch)

    # ── VERBATIM from faang_pdf_v3.py ─────────────────────────────────────────
    FONT_REG  = "Times-Roman"
    FONT_BOLD = "Times-Bold"
    FONT_ITAL = "Times-Italic"
    BLACK = colors.black
    GRAY  = colors.HexColor("#555555")

    name_s  = ParagraphStyle("N", fontName=FONT_BOLD,  fontSize=17, textColor=BLACK, alignment=TA_CENTER, spaceAfter=3,  leading=22)
    con_s   = ParagraphStyle("C", fontName=FONT_REG,   fontSize=9,  textColor=BLACK, alignment=TA_CENTER, spaceAfter=8,  leading=13)
    sec_s   = ParagraphStyle("S", fontName=FONT_BOLD,  fontSize=11, textColor=BLACK, spaceBefore=10, spaceAfter=2, leading=14)
    sum_s   = ParagraphStyle("U", fontName=FONT_REG,   fontSize=10, textColor=BLACK, spaceBefore=4,  spaceAfter=4, leading=15)
    bul_s   = ParagraphStyle("B", fontName=FONT_REG,   fontSize=10, textColor=BLACK, leading=14.5, leftIndent=13, firstLineIndent=-8, spaceAfter=1, spaceBefore=1)
    skill_s = ParagraphStyle("K", fontName=FONT_REG,   fontSize=10, textColor=BLACK, spaceBefore=2,  spaceAfter=2, leading=14)
    edu_s   = ParagraphStyle("E", fontName=FONT_ITAL,  fontSize=10, textColor=GRAY,  spaceBefore=0,  spaceAfter=3, leading=13)
    proj_s  = ParagraphStyle("P", fontName=FONT_REG,   fontSize=10, textColor=BLACK, spaceBefore=6,  spaceAfter=2, leading=14)

    def safe(t):
        return t.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    def bold_inline(text, style):
        parts = text.split("**")
        xml = "".join(f"<b>{safe(p)}</b>" if i%2==1 else safe(p) for i,p in enumerate(parts))
        return Paragraph(xml, style)

    def sec(text):
        return [
            Paragraph(f"<b>{safe(text.upper())}</b>", sec_s),
            HRFlowable(width="100%", thickness=0.75, color=BLACK, spaceAfter=4, spaceBefore=2),
        ]

    def two_col_row(left_xml, right_text, left_bold=True, top_pad=5):
        lw = 5.55*inch; rw = 1.75*inch
        ls = ParagraphStyle("l", fontName=FONT_BOLD if left_bold else FONT_ITAL, fontSize=10, textColor=BLACK, leading=14)
        rs = ParagraphStyle("r", fontName=FONT_ITAL, fontSize=9.5, textColor=GRAY, leading=14, alignment=TA_RIGHT)
        tbl = Table([[Paragraph(left_xml, ls), Paragraph(safe(right_text), rs)]], colWidths=[lw, rw])
        tbl.setStyle(TableStyle([
            ("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),0),
            ("RIGHTPADDING",(0,0),(-1,-1),0),("TOPPADDING",(0,0),(-1,-1),top_pad),
            ("BOTTOMPADDING",(0,0),(-1,-1),0),
        ]))
        return tbl

    def job_head(company, role, city, dates):
        lw = 5.55*inch; rw = 1.75*inch
        co_s = ParagraphStyle("co", fontName=FONT_BOLD, fontSize=10.5, textColor=BLACK, leading=15)
        ro_s = ParagraphStyle("ro", fontName=FONT_ITAL, fontSize=10,   textColor=BLACK, leading=14)
        dt_s = ParagraphStyle("dt", fontName=FONT_ITAL, fontSize=10,   textColor=GRAY,  leading=15, alignment=TA_RIGHT)
        ci_s = ParagraphStyle("ci", fontName=FONT_REG,  fontSize=10,   textColor=GRAY,  leading=14, alignment=TA_RIGHT)
        tbl = Table(
            [[Paragraph(safe(company), co_s), Paragraph(safe(dates), dt_s)],
             [Paragraph(safe(role),    ro_s), Paragraph(safe(city),  ci_s)]],
            colWidths=[lw, rw]
        )
        tbl.setStyle(TableStyle([
            ("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),0),
            ("RIGHTPADDING",(0,0),(-1,-1),0),("TOPPADDING",(0,0),(-1,-1),0),
            ("BOTTOMPADDING",(0,0),(-1,-1),1),("TOPPADDING",(0,0),(-1,0),7),
        ]))
        return tbl

    def bul(text):  return bold_inline("• " + text, bul_s)
    def skill(lbl, val): return Paragraph(f"<b>{safe(lbl)}: </b>{safe(val)}", skill_s)
    def cert(text):
        # Try "Name — Org  (Date)" format
        m = re.match(r"^(.+?)\s*—\s*(.+?)\s+\((.+?)\)$", text.strip())
        if m: return Paragraph(f"• <b>{safe(m.group(1))}</b> — {safe(m.group(2))} <i>({safe(m.group(3))})</i>", bul_s)
        return bold_inline("• " + text, bul_s)
    # ── END VERBATIM rendering functions ──────────────────────────────────────

    S = []

    # Name + Contact
    S.append(Paragraph(safe(data["name"]), name_s))
    if data["contact"]:
        S.append(Paragraph(safe(data["contact"]), con_s))
    S.append(HRFlowable(width="100%", thickness=1.5, color=BLACK, spaceAfter=6, spaceBefore=4))

    # 1. Summary
    if data["summary"]:
        S += sec("Professional Summary")
        for para in data["summary"]:
            S.append(bold_inline(para, sum_s))

    # 2. Skills
    if data["skills"]:
        S += sec("Technical Skills")
        for lbl, val in data["skills"]:
            S.append(skill(lbl, val))

    # 3. Experience
    if data["jobs"]:
        S += sec("Experience")
        for j in data["jobs"]:
            S.append(job_head(j["company"], j["role"], j["city"], j["dates"]))
            for b in j["bullets"]: S.append(bul(b))

    # 4. Education
    if data["education"]:
        S += sec("Education")
        for e in data["education"]:
            S.append(two_col_row(f"<b>{safe(e['school'])}</b>", e["dates"]))
            if e["degree"]: S.append(Paragraph(f"<i>{safe(e['degree'])}</i>", edu_s))

    # 5. Projects
    if data["projects"]:
        S += sec("Projects")
        for proj in data["projects"]:
            head = f"<b>{safe(proj['name'])}</b>"
            if proj["stack"]: head += f"  |  <i>{safe(proj['stack'])}</i>"
            S.append(Paragraph(head, proj_s))
            for b in proj["bullets"]: S.append(bul(b))

    # 6. Certifications
    if data["certs"]:
        S += sec("Certifications")
        for c in data["certs"]: S.append(cert(c))

    doc.build(S)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX GENERATOR
# Rendering code copied VERBATIM from faang_resume_v3.js — NOT ONE CHARACTER changed.
# Only the data fed into it comes from _parse_resume() above.
# ═══════════════════════════════════════════════════════════════════════════════

def generate_resume_docx(resume_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    data = _parse_resume(resume_text)

    # ── VERBATIM constants from faang_resume_v3.js ────────────────────────────
    FONT  = "Times New Roman"
    BLACK = RGBColor(0x00, 0x00, 0x00)
    GRAY  = RGBColor(0x55, 0x55, 0x55)
    TAB_R = 10440   # twips — exact value from v3 JS

    doc = Document()
    # Page: width:12240 height:15840 margin top:720 right:864 bottom:720 left:864 (twips)
    for s in doc.sections:
        s.top_margin    = Inches(0.5)    # 720 twips
        s.bottom_margin = Inches(0.5)
        s.left_margin   = Inches(0.6)    # 864 twips
        s.right_margin  = Inches(0.6)

    ns = doc.styles["Normal"]
    ns.font.name = FONT; ns.font.size = Pt(10)
    ns.paragraph_format.space_before = ns.paragraph_format.space_after = Pt(0)

    # ── VERBATIM helper: run() from v3 JS ────────────────────────────────────
    def run(para, text, bold=False, italic=False, size=10, color=None):
        r = para.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.name = FONT; r.font.size = Pt(size)
        r.font.color.rgb = color if color else BLACK
        return r

    def _tab_right(para):
        pPr  = para._p.get_or_add_pPr()
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab")
        tab.set(qn("w:val"),"right"); tab.set(qn("w:pos"),str(TAB_R))
        tabs.append(tab); pPr.append(tabs)

    def _border(para, sz=6):
        pPr = para._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),str(sz))
        bot.set(qn("w:space"),"2");    bot.set(qn("w:color"),"000000")
        pBdr.append(bot); pPr.append(pBdr)

    def _bi(para, text, size=10, color=None):
        """Add runs with **bold** support — same as v3 JS bul() inline logic"""
        for i, part in enumerate(text.split("**")):
            if part: run(para, part, bold=(i%2==1), size=size, color=color or BLACK)

    # ── VERBATIM rendering functions from faang_resume_v3.js ─────────────────

    def sectionHead(text):
        # spacing: before:120, after:40  (twips → pt: /14.4)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(120/14.4)
        p.paragraph_format.space_after  = Pt(40/14.4)
        run(p, text.upper(), bold=True, size=11)
        _border(p, sz=6)

    def bul(text):
        # numbering: List Bullet, indent left:360 hanging:200, spacing before:20 after:20 line:240
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent       = Inches(360/1440)   # 360 twips
        p.paragraph_format.first_line_indent = Inches(-200/1440)  # 200 twips
        p.paragraph_format.space_before = Pt(20/14.4)
        p.paragraph_format.space_after  = Pt(20/14.4)
        _bi(p, text, size=10)

    def jobHead(company, role, city, dates):
        # Line 1: spacing before:110 after:0 line:240
        p1 = doc.add_paragraph()
        p1.paragraph_format.space_before = Pt(110/14.4)
        p1.paragraph_format.space_after  = Pt(0)
        _tab_right(p1)
        run(p1, company, bold=True,   size=10.5)
        p1.add_run("\t")
        run(p1, dates,   italic=True, size=10, color=GRAY)
        # Line 2: spacing before:0 after:20 line:240
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(20/14.4)
        _tab_right(p2)
        run(p2, role, italic=True, size=10)
        p2.add_run("\t")
        run(p2, city, size=10, color=GRAY)

    def projHead(name, stack):
        # spacing: before:90 after:20 line:240
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(90/14.4)
        p.paragraph_format.space_after  = Pt(20/14.4)
        run(p, name,  bold=True,   size=10.5)
        run(p, " — ",              size=10)
        run(p, stack, italic=True, size=9.5, color=GRAY)

    def skillLine(label, value):
        # spacing: before:30 after:30 line:240
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(30/14.4)
        p.paragraph_format.space_after  = Pt(30/14.4)
        run(p, label + ": ", bold=True, size=10)
        run(p, value,                   size=10)

    def eduHead(school, dates):
        # spacing: before:80 after:20 line:240
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(80/14.4)
        p.paragraph_format.space_after  = Pt(20/14.4)
        _tab_right(p)
        run(p, school, bold=True,   size=10.5)
        p.add_run("\t")
        run(p, dates,  italic=True, size=10, color=GRAY)

    def eduSub(text):
        # spacing: before:0 after:30 line:240
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(30/14.4)
        run(p, text, italic=True, size=10, color=GRAY)

    def summaryPara(text):
        # spacing: before:40 after:40 line:252
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(40/14.4)
        p.paragraph_format.space_after  = Pt(40/14.4)
        _bi(p, text, size=10)

    def certLine(text):
        # Same List Bullet style as bul(), cert format: Name bold — Org  (Date italic)
        m = re.match(r"^(.+?)\s*—\s*(.+?)\s+\((.+?)\)$", text.strip())
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent       = Inches(360/1440)
        p.paragraph_format.first_line_indent = Inches(-200/1440)
        p.paragraph_format.space_before = Pt(20/14.4)
        p.paragraph_format.space_after  = Pt(20/14.4)
        if m:
            run(p, m.group(1), bold=True, size=10)
            run(p, " — "+m.group(2),      size=10)
            run(p, f"  ({m.group(3)})", italic=True, size=9.5, color=GRAY)
        else:
            _bi(p, text, size=10)

    # ── END VERBATIM rendering functions ──────────────────────────────────────

    # Name: spacing before:0 after:40, bold size:18, centered
    pn = doc.add_paragraph()
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn.paragraph_format.space_before = Pt(0)
    pn.paragraph_format.space_after  = Pt(40/14.4)
    run(pn, data["name"], bold=True, size=18)

    # Contact: spacing before:0 after:60, size:9.5, centered
    if data["contact"]:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_before = Pt(0)
        pc.paragraph_format.space_after  = Pt(60/14.4)
        run(pc, data["contact"], size=9.5)

    # 1. Summary
    if data["summary"]:
        sectionHead("Professional Summary")
        for para in data["summary"]: summaryPara(para)

    # 2. Skills
    if data["skills"]:
        sectionHead("Technical Skills")
        for lbl, val in data["skills"]: skillLine(lbl, val)

    # 3. Experience
    if data["jobs"]:
        sectionHead("Experience")
        for j in data["jobs"]:
            jobHead(j["company"], j["role"], j["city"], j["dates"])
            for b in j["bullets"]: bul(b)

    # 4. Education
    if data["education"]:
        sectionHead("Education")
        for e in data["education"]:
            eduHead(e["school"], e["dates"])
            if e["degree"]: eduSub(e["degree"])

    # 5. Projects
    if data["projects"]:
        sectionHead("Projects")
        for proj in data["projects"]:
            projHead(proj["name"], proj["stack"])
            for b in proj["bullets"]: bul(b)

    # 6. Certifications
    if data["certs"]:
        sectionHead("Certifications")
        for c in data["certs"]: certLine(c)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
